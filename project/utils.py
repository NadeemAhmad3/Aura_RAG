import os
import shutil
import config
import pymupdf.layout
import pymupdf4llm
from pathlib import Path
import glob
import tiktoken


def clear_directory_contents(directory: Path) -> None:
    """Delete everything under directory but not the directory itself (safe for Docker volume / bind mount roots)."""
    directory = Path(directory)
    if not directory.is_dir():
        return
    for child in directory.iterdir():
        if child.is_dir():
            shutil.rmtree(child)
        else:
            child.unlink()


os.environ["TOKENIZERS_PARALLELISM"] = "false"

def pdf_to_markdown(pdf_path, output_dir):
    doc = pymupdf.open(pdf_path)
    md = pymupdf4llm.to_markdown(doc, header=False, footer=False, page_separators=True, ignore_images=True, write_images=False, image_path=None)
    md_cleaned = md.encode('utf-8', errors='surrogatepass').decode('utf-8', errors='ignore')
    output_path = Path(output_dir) / Path(doc.name).stem
    Path(output_path).with_suffix(".md").write_bytes(md_cleaned.encode('utf-8'))

def pdfs_to_markdowns(path_pattern, overwrite: bool = False):
    output_dir = Path(config.MARKDOWN_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    for pdf_path in map(Path, glob.glob(path_pattern)):
        md_path = (output_dir / pdf_path.stem).with_suffix(".md")
        if overwrite or not md_path.exists():
            pdf_to_markdown(pdf_path, output_dir)

def estimate_context_tokens(messages: list) -> int:
    try:
        encoding = tiktoken.encoding_for_model("gpt-4")
    except:
        encoding = tiktoken.get_encoding("cl100k_base")
    return sum(len(encoding.encode(str(msg.content))) for msg in messages if hasattr(msg, 'content') and msg.content)


def docx_to_markdown(docx_path, output_dir):
    import docx
    doc = docx.Document(docx_path)
    md_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
        if para.style.name.startswith("Heading 1"):
            md_lines.append(f"# {text}")
        elif para.style.name.startswith("Heading 2"):
            md_lines.append(f"## {text}")
        elif para.style.name.startswith("Heading 3"):
            md_lines.append(f"### {text}")
        elif para.style.name.startswith("List Bullet"):
            md_lines.append(f"- {text}")
        elif para.style.name.startswith("List Number"):
            md_lines.append(f"1. {text}")
        else:
            md_lines.append(text)
            
    for table in doc.tables:
        md_lines.append("")
        for i, row in enumerate(table.rows):
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_lines.append("| " + " | ".join(row_cells) + " |")
            if i == 0:
                md_lines.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
        md_lines.append("")
        
    doc_name = Path(docx_path).name.replace(".docx", "")
    output_path = Path(output_dir) / f"{doc_name}.md"
    output_path.write_text("\n".join(md_lines), encoding="utf-8")


def csv_to_markdown(csv_path, output_dir):
    import csv
    md_lines = []
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
    except UnicodeDecodeError:
        with open(csv_path, 'r', encoding='latin-1') as f:
            reader = csv.reader(f)
            rows = list(reader)
            
    if rows:
        headers = [h.strip().replace("\n", " ") for h in rows[0]]
        md_lines.append("| " + " | ".join(headers) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows[1:]:
            cells = [cell.strip().replace("\n", " ") for cell in row]
            if len(cells) < len(headers):
                cells += [""] * (len(headers) - len(cells))
            elif len(cells) > len(headers):
                cells = cells[:len(headers)]
            md_lines.append("| " + " | ".join(cells) + " |")
            
    doc_name = Path(csv_path).name.replace(".csv", "")
    output_path = Path(output_dir) / f"{doc_name}.md"
    output_path.write_text("\n".join(md_lines), encoding="utf-8")


def txt_to_markdown(txt_path, output_dir):
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(txt_path, 'r', encoding='latin-1') as f:
            content = f.read()
            
    doc_name = Path(txt_path).name.replace(".txt", "")
    output_path = Path(output_dir) / f"{doc_name}.md"
    output_path.write_text(content, encoding="utf-8")


def url_to_markdown(url, output_dir):
    import requests
    from bs4 import BeautifulSoup, Comment
    import re
    from pathlib import Path
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    
    def _fetch_github_contributions(username):
        ignored_names = {"search", "explore", "trending", "features", "pricing", "security", "customer-stories", "enterprise", "readme", "about", "contact", "join", "login", "signup"}
        if username.lower() not in ignored_names:
            try:
                contrib_url = f"https://github.com/users/{username}/contributions"
                contrib_resp = requests.get(contrib_url, headers=headers, timeout=10)
                if contrib_resp.status_code == 200:
                    contrib_soup = BeautifulSoup(contrib_resp.text, 'html.parser')
                    contrib_header = contrib_soup.find('h2')
                    if contrib_header:
                        return " ".join(contrib_header.text.split())
            except Exception as e:
                print(f"Error fetching contributions for {username}: {e}")
        return None

    # 1. Try RAG-tuned markdown converter API (Jina Reader) first (handles JS, tables, boilerplate)
    try:
        jina_url = f"https://r.jina.ai/{url}"
        response = requests.get(jina_url, headers=headers, timeout=12)
        if response.status_code == 200 and response.text.strip():
            safe_name = re.sub(r'[^a-zA-Z0-9]', '_', url.replace("https://", "").replace("http://", ""))[:50]
            output_path = Path(output_dir) / f"{safe_name}.md"
            
            content = response.text
            github_profile_match = re.match(r"https?://(?:www\.)?github\.com/([^/]+)/?$", url)
            if github_profile_match:
                username = github_profile_match.group(1)
                contrib_text = _fetch_github_contributions(username)
                if contrib_text:
                    content = f"# GitHub Profile: {username}\nSource: {url}\n\n## GitHub Contributions Overview\n**{contrib_text}**\n\n" + content
            
            output_path.write_text(content, encoding="utf-8")
            return safe_name
    except Exception as e:
        print(f"Jina Reader fallback to local scraper due to: {e}")

    # 2. Local Semantic DOM Scraper Fallback
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()
    
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Remove HTML Comments
    for comment in soup.find_all(text=lambda text: isinstance(text, Comment)):
        comment.extract()
        
    # Remove standard noise blocks (ads, cookies, sidebars, share widgets)
    noise_selectors = [
        "script", "style", "iframe", "noscript", "svg", "canvas",
        "header", "footer", "nav", "aside",
        ".header", ".footer", ".nav", ".navigation", ".menu", ".sidebar",
        ".cookie", ".consent", ".banner", ".ads", ".advertisement", ".social-share",
        "#header", "#footer", "#nav", "#navigation", "#menu", "#sidebar", "#cookie"
    ]
    
    for selector in noise_selectors:
        for element in soup.select(selector):
            element.decompose()
            
    # Recursive parser to convert HTML elements to structured Markdown
    def html_to_markdown(element):
        if element.name is None:
            return element.strip()
            
        child_text = []
        for child in element.children:
            txt = html_to_markdown(child)
            if txt:
                child_text.append(txt)
        inner_content = " ".join(child_text).strip()
        
        name = element.name.lower()
        
        if name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(name[1])
            return f"\n\n" + ("#" * level) + f" {inner_content}\n\n"
        elif name == 'p':
            return f"\n\n{inner_content}\n\n"
        elif name == 'br':
            return "\n"
        elif name == 'li':
            parent = element.parent
            if parent and parent.name == 'ol':
                siblings = [c for c in parent.children if c.name == 'li']
                try:
                    idx = siblings.index(element) + 1
                except ValueError:
                    idx = 1
                return f"\n{idx}. {inner_content}"
            return f"\n- {inner_content}"
        elif name in ['ul', 'ol']:
            return f"\n{inner_content}\n"
        elif name == 'blockquote':
            lines = [f"> {line}" for line in inner_content.split('\n')]
            return f"\n\n" + "\n".join(lines) + f"\n\n"
        elif name in ['pre', 'code']:
            if name == 'pre' or '\n' in inner_content:
                return f"\n\n```\n{inner_content}\n```\n\n"
            return f"`{inner_content}`"
        elif name == 'a':
            href = element.get('href', '')
            if href and not href.startswith('javascript:') and not href.startswith('#'):
                return f" [{inner_content}]({href}) "
            return inner_content
        elif name == 'table':
            table_lines = []
            rows = element.find_all('tr')
            for r_idx, row in enumerate(rows):
                cols = row.find_all(['th', 'td'])
                cells = [col.get_text().strip().replace("\n", " ") for col in cols]
                if cells:
                    table_lines.append("| " + " | ".join(cells) + " |")
                    if r_idx == 0:
                        table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            return f"\n\n" + "\n".join(table_lines) + f"\n\n"
        elif name in ['div', 'span', 'article', 'section', 'main', 'body']:
            return inner_content
            
        return inner_content

    md_content = []
    title = soup.title.string.strip() if soup.title else ""
    if title:
        md_content.append(f"# {title}\n")
        md_content.append(f"Source: {url}\n\n")

    # Handle custom case for GitHub Profile fallback locally (for calendar graph text)
    github_profile_match = re.match(r"https?://(?:www\.)?github\.com/([^/]+)/?$", url)
    if github_profile_match:
        username = github_profile_match.group(1)
        contrib_text = _fetch_github_contributions(username)
        if contrib_text:
            md_content.append(f"## GitHub Contributions Overview\n")
            md_content.append(f"**{contrib_text}**\n\n")

    # Process core document content
    body = soup.find('article') or soup.find('main') or soup.find('body') or soup
    body_markdown = html_to_markdown(body)
    md_content.append(body_markdown)
    
    content = "\n".join(md_content)
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    safe_name = re.sub(r'[^a-zA-Z0-9]', '_', url.replace("https://", "").replace("http://", ""))[:50]
    output_path = Path(output_dir) / f"{safe_name}.md"
    output_path.write_text(content, encoding="utf-8")
    return safe_name
